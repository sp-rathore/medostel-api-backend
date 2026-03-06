#!/usr/bin/env python3
"""
MOSL AI Platform Strategy 2026 - Word to HTML Converter
Converts a professionally formatted Word document to a beautifully styled HTML page.
"""

import os
import sys
import base64
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell
from PIL import Image
import io

# Configuration
SOURCE_DOCX = "/Users/shishupals/Documents/Claude/projects/Medostel/repositories/medostel-api-backend/prompt/MOSL AI Platform Strategy 2026.docx"
OUTPUT_HTML = "/Users/shishupals/Documents/Claude/projects/Medostel/repositories/medostel-api-backend/prompt/MOSL_AI_Platform_Strategy_2026.html"

# Color scheme
HEADER_COLOR = "#1E40AF"  # Professional Navy Blue
TABLE_HEADER_BG = "#EFF6FF"  # Light navy blue
TEXT_COLOR = "#000000"  # Black
BORDER_COLOR = "#CBD5E1"  # Light gray

class DocxExtractor:
    def __init__(self, docx_path: str):
        self.doc = Document(docx_path)
        self.elements = []
        self.images = {}
        self.image_count = 0
        self.image_paragraphs = []  # Track which paragraphs have images

    def extract_all(self) -> Dict:
        """Extract all content from document"""
        print("Extracting document content...")

        # First pass: extract all images from inline shapes
        self._extract_inline_shapes()

        # Second pass: process paragraphs and tables
        para_index = 0
        for element in self.doc.element.body:
            if element.tag.endswith('p'):
                para = element
                self._process_paragraph(para, para_index)
                para_index += 1
            elif element.tag.endswith('tbl'):
                self._process_table(element)

        return {
            'elements': self.elements,
            'images': self.images,
            'image_count': self.image_count,
            'image_paragraphs': self.image_paragraphs
        }

    def _extract_inline_shapes(self):
        """Extract all images from inline shapes"""
        print("Extracting images from document...")
        try:
            for shape in self.doc.part.inline_shapes:
                try:
                    # Get the image from the inline shape
                    image_part = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
                    # Access the actual image part
                    image_rel = self.doc.part.rels.get(image_part)
                    if image_rel:
                        image_part_obj = image_rel.target_part
                        image_blob = image_part_obj.blob
                        image_format = self._detect_image_format(image_blob)

                        # Convert WDP to PNG if needed
                        if image_format == 'wdp':
                            image_blob = self._convert_wdp_to_png(image_blob)
                            image_format = 'png'

                        # Store image as base64
                        image_b64 = base64.b64encode(image_blob).decode('utf-8')
                        image_id = f"image_{self.image_count}"
                        self.image_count += 1

                        self.images[image_id] = {
                            'base64': image_b64,
                            'format': image_format
                        }

                        print(f"  Extracted {image_id} ({image_format})")
                except Exception as e:
                    print(f"  Warning: Could not extract inline shape: {e}")
        except Exception as e:
            print(f"  Warning: Error extracting images: {e}")

    def _process_paragraph(self, para_element, para_index):
        """Process a paragraph element"""
        para = None
        # Find matching paragraph in doc.paragraphs
        for p in self.doc.paragraphs:
            if p._element == para_element:
                para = p
                break

        if para is None:
            return

        text = para.text
        style_name = para.style.name if para.style else 'Normal'

        # Check for images in this paragraph
        has_image = bool(para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
        if has_image:
            self.image_paragraphs.append(para_index)

        element_data = {
            'type': 'paragraph',
            'style': style_name,
            'text': text,
            'runs': self._extract_runs(para),
            'has_image': has_image,
            'para_index': para_index
        }

        self.elements.append(element_data)

    def _process_table(self, table_element):
        """Process a table element"""
        # Find matching table in doc.tables
        table = None
        for tbl in self.doc.tables:
            if tbl._element == table_element:
                table = tbl
                break

        if table is None:
            return

        rows_data = []
        for i, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                cell_data = {
                    'text': cell.text,
                    'paragraphs': [p.text for p in cell.paragraphs]
                }
                row_data.append(cell_data)
            rows_data.append(row_data)

        element_data = {
            'type': 'table',
            'rows': rows_data,
            'row_count': len(table.rows),
            'col_count': len(table.rows[0].cells) if table.rows else 0
        }

        self.elements.append(element_data)

    def _extract_runs(self, para) -> List[Dict]:
        """Extract run-level formatting"""
        runs_data = []
        for run in para.runs:
            run_data = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_size': run.font.size.pt if run.font.size else None,
                'font_name': run.font.name,
                'color': self._get_run_color(run)
            }
            runs_data.append(run_data)
        return runs_data

    def _get_run_color(self, run) -> Optional[str]:
        """Extract color from run"""
        try:
            color = run.font.color.rgb
            if color:
                return str(color)
        except:
            pass
        return None

    def _detect_image_format(self, blob: bytes) -> str:
        """Detect image format from bytes"""
        if blob.startswith(b'\x89PNG'):
            return 'png'
        elif blob.startswith(b'\xff\xd8\xff'):
            return 'jpg'
        elif blob.startswith(b'II') or blob.startswith(b'MM'):
            return 'wdp'
        return 'unknown'

    def _convert_wdp_to_png(self, wdp_blob: bytes) -> bytes:
        """Convert WDP to PNG"""
        try:
            img = Image.open(io.BytesIO(wdp_blob))
            png_blob = io.BytesIO()
            img.save(png_blob, format='PNG')
            return png_blob.getvalue()
        except Exception as e:
            print(f"Warning: Could not convert WDP: {e}")
            return wdp_blob


class HTMLGenerator:
    def __init__(self, elements: List[Dict], images: Dict, image_paragraphs: List[int] = None):
        self.elements = elements
        self.images = images
        self.image_paragraphs = image_paragraphs or []

    def generate(self) -> str:
        """Generate complete HTML document"""
        print("Generating HTML...")

        css = self._generate_css()
        body = self._generate_body()

        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>MOSL AI Platform Strategy 2026</title>
    <style>
{css}
    </style>
</head>
<body>
{body}
</body>
</html>"""

        return html

    def _generate_css(self) -> str:
        """Generate CSS styling"""
        return """        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: """ + TEXT_COLOR + """;
            background-color: #f8f9fa;
            padding: 40px 20px;
        }

        .container {
            max-width: 1000px;
            margin: 0 auto;
            background-color: white;
            padding: 60px;
            box-shadow: 0 2px 10px rgba(0, 0, 0, 0.1);
            border-radius: 4px;
        }

        h1, h2, h3, h4 {
            color: """ + HEADER_COLOR + """;
            margin-top: 30px;
            margin-bottom: 15px;
            font-weight: bold;
        }

        h1 {
            font-size: 32px;
            border-bottom: 3px solid """ + HEADER_COLOR + """;
            padding-bottom: 10px;
            margin-top: 40px;
        }

        h2 {
            font-size: 28px;
            margin-top: 35px;
        }

        h3 {
            font-size: 24px;
            margin-top: 25px;
        }

        h4 {
            font-size: 20px;
            margin-top: 20px;
        }

        p {
            margin-bottom: 15px;
            text-align: justify;
        }

        p.list-paragraph {
            margin-left: 30px;
            margin-bottom: 10px;
        }

        table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
            box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
        }

        th {
            background-color: """ + TABLE_HEADER_BG + """;
            color: """ + TEXT_COLOR + """;
            padding: 15px;
            text-align: left;
            font-weight: bold;
            border: 1px solid """ + BORDER_COLOR + """;
        }

        td {
            padding: 12px 15px;
            border: 1px solid """ + BORDER_COLOR + """;
            background-color: white;
        }

        tr:nth-child(even) td {
            background-color: #f9fafc;
        }

        tr:hover td {
            background-color: #f0f4f8;
        }

        img {
            max-width: 100%;
            height: auto;
            margin: 20px 0;
            border-radius: 4px;
            box-shadow: 0 2px 8px rgba(0, 0, 0, 0.1);
        }

        ul, ol {
            margin: 15px 0 15px 40px;
        }

        li {
            margin-bottom: 8px;
        }

        strong {
            font-weight: bold;
        }

        em {
            font-style: italic;
        }

        u {
            text-decoration: underline;
        }

        @media (max-width: 768px) {
            .container {
                padding: 30px;
            }

            h1 { font-size: 28px; }
            h2 { font-size: 24px; }
            h3 { font-size: 20px; }
            h4 { font-size: 18px; }

            table {
                font-size: 14px;
            }

            th, td {
                padding: 10px;
            }
        }"""

    def _generate_body(self) -> str:
        """Generate HTML body content"""
        body_html = '    <div class="container">\n'

        element_count = 0
        image_index = 0

        for element in self.elements:
            element_count += 1

            if element['type'] == 'paragraph':
                body_html += self._generate_paragraph(element)
                # Add image if this paragraph has one
                if element.get('has_image') and image_index < len(self.images):
                    image_id = f"image_{image_index}"
                    if image_id in self.images:
                        body_html += self._generate_image(image_id, self.images[image_id])
                    image_index += 1
            elif element['type'] == 'table':
                body_html += self._generate_table(element)

        body_html += '    </div>\n'

        print(f"  Generated {element_count} elements")

        return body_html

    def _generate_paragraph(self, para: Dict) -> str:
        """Generate paragraph HTML"""
        style = para['style']
        text = para['text']

        # Determine heading level
        if style.startswith('Heading'):
            level = style.replace('Heading', '').strip()
            if level in ['1', '2', '3', '4']:
                tag = f"h{level}"
                return f'        <{tag}>{self._escape_html(text)}</{tag}>\n'

        # Check if it's a list paragraph
        if style == 'List Paragraph' or 'List' in style:
            html = f'        <p class="list-paragraph">• {self._escape_html(text)}</p>\n'
        else:
            # Regular paragraph
            html = '        <p>'

            # Apply run-level formatting
            if para.get('runs'):
                for run in para['runs']:
                    run_html = self._escape_html(run['text'])
                    if run['bold']:
                        run_html = f"<strong>{run_html}</strong>"
                    if run['italic']:
                        run_html = f"<em>{run_html}</em>"
                    if run['underline']:
                        run_html = f"<u>{run_html}</u>"
                    html += run_html
            else:
                html += self._escape_html(text)

            html += '</p>\n'

        return html

    def _generate_image(self, img_id: str, img_data: Dict) -> str:
        """Generate image HTML"""
        image_b64 = img_data['base64']
        image_format = img_data['format']
        mime_type = f"image/{image_format}"
        if image_format == 'jpg':
            mime_type = 'image/jpeg'

        return f'        <img src="data:{mime_type};base64,{image_b64}" alt="Document image">\n'

    def _generate_table(self, table: Dict) -> str:
        """Generate table HTML"""
        rows = table['rows']
        if not rows:
            return ''

        html = '        <table>\n'

        # Header row
        html += '            <thead>\n                <tr>\n'
        for cell in rows[0]:
            html += f"                    <th>{self._escape_html(cell['text'])}</th>\n"
        html += '                </tr>\n            </thead>\n'

        # Body rows
        html += '            <tbody>\n'
        for row in rows[1:]:
            html += '                <tr>\n'
            for cell in row:
                html += f"                    <td>{self._escape_html(cell['text'])}</td>\n"
            html += '                </tr>\n'
        html += '            </tbody>\n'

        html += '        </table>\n'
        return html

    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters"""
        return (text
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;')
            .replace("'", '&#39;')
        )


def main():
    print("=" * 60)
    print("MOSL AI Platform Strategy 2026 - Word to HTML Converter")
    print("=" * 60)

    # Check if source file exists
    if not os.path.exists(SOURCE_DOCX):
        print(f"ERROR: Source file not found: {SOURCE_DOCX}")
        sys.exit(1)

    print(f"\nSource: {SOURCE_DOCX}")
    print(f"Output: {OUTPUT_HTML}\n")

    # Extract content
    extractor = DocxExtractor(SOURCE_DOCX)
    data = extractor.extract_all()

    print(f"\nExtraction complete:")
    print(f"  - Total elements: {len(data['elements'])}")
    print(f"  - Images extracted: {data['image_count']}")

    # Generate HTML
    generator = HTMLGenerator(data['elements'], data['images'], data['image_paragraphs'])
    html_content = generator.generate()

    # Count validation data
    para_count = sum(1 for e in data['elements'] if e['type'] == 'paragraph')
    table_count = sum(1 for e in data['elements'] if e['type'] == 'table')

    # Write output
    with open(OUTPUT_HTML, 'w', encoding='utf-8') as f:
        f.write(html_content)

    print(f"\n✓ HTML file generated successfully!")
    print(f"  Location: {OUTPUT_HTML}")
    print(f"  Size: {len(html_content) / 1024:.1f} KB")

    # Validation
    print("\nValidation:")
    print(f"  ✓ Paragraphs: {para_count}")
    print(f"  ✓ Tables: {table_count}")
    print(f"  ✓ Images: {data['image_count']}")

    print("\n" + "=" * 60)
    print("Conversion complete! Open the HTML file in a web browser.")
    print("=" * 60)


if __name__ == '__main__':
    main()
